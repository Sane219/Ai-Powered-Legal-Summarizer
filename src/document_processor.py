"""
Document processing module for legal documents
Handles PDF, DOCX, and text file extraction and preprocessing
"""
import re
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import logging
from typing import Dict, List, Optional, Tuple
import spacy
from spacy import displacy

from config import SUPPORTED_FORMATS, LEGAL_SECTIONS, CLAUSE_KEYWORDS

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles extraction and preprocessing of legal documents"""
    
    def __init__(self):
        """Initialize the document processor with spaCy model"""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            try:
                # Try to download the model automatically
                import subprocess
                import sys
                subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
                self.nlp = spacy.load("en_core_web_sm")
                logger.info("Successfully downloaded and loaded spaCy model")
            except Exception as e:
                logger.warning(f"Could not load or download spaCy model: {e}")
                self.nlp = None
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text() + "\n"
            
            doc.close()
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file: {e}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from supported file formats"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() not in SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        if file_path.suffix.lower() == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            return self.extract_text_from_docx(str(file_path))
        elif file_path.suffix.lower() == '.txt':
            return self.extract_text_from_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    def clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)
        
        return text.strip()
    
    def identify_legal_sections(self, text: str) -> Dict[str, List[str]]:
        """Identify and extract legal sections from the document"""
        sections = {section: [] for section in LEGAL_SECTIONS}
        
        # Split text into paragraphs
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        for section_name, keywords in CLAUSE_KEYWORDS.items():
            for paragraph in paragraphs:
                # Check if paragraph contains keywords for this section
                if any(keyword.lower() in paragraph.lower() for keyword in keywords):
                    if section_name in sections:
                        sections[section_name].append(paragraph)
        
        return sections
    
    def extract_entities(self, text: str) -> Dict[str, List[Dict]]:
        """Extract named entities from legal text using spaCy or fallback method"""
        if not self.nlp:
            # Fallback entity extraction using regex patterns
            return self._extract_entities_fallback(text)
        
        try:
            doc = self.nlp(text)
            entities = []
            
            for ent in doc.ents:
                entities.append({
                    "text": ent.text,
                    "label": ent.label_,
                    "start": ent.start_char,
                    "end": ent.end_char,
                    "description": spacy.explain(ent.label_) if spacy.explain(ent.label_) else ent.label_
                })
            
            # Group entities by type
            entity_groups = {}
            for entity in entities:
                label = entity["label"]
                if label not in entity_groups:
                    entity_groups[label] = []
                entity_groups[label].append(entity)
            
            return {
                "entities": entities,
                "groups": entity_groups,
                "total_count": len(entities)
            }
        except Exception as e:
            logger.warning(f"spaCy entity extlback")
    ext)
    
    def _extract_entities_fallback(self, text: str) -> Dict[str, Li:
        """Fallback entit
        entities = []
        
        # Basic patterns for common entities
        patterns = {
            "ORG": [
                r'\b[A-Z][a-zA-Z\s&',
         )\b'
         ],
            "PERSO": [
                r'\b[A-Z][a-z]+\s+[A-',
                r'\bMr\.?\s+[A-Z][a-z]+\b',
                r'\bMs\.?\s+[A-Z]]+\b',
                r'\bDr\.?\s+[A-Z][a-z]+\b'
            ],
            "DATE": [
                r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{',
                
            ],
            "MONEY": [
                r'\$[\d,]+(?:\.\d{2})?\
                r'\b\d+\s+dollars?\b'
            ]
        
        
    tems():
            for pattern in pattern_list:
                matches = re.finditer(pattern, text, re.IGNOECASE)
                for match in matches:
                    entiti({
                        "text": match.group(),
                        "label": label,
                        "start": match.start(),
                        "end": match.end(),
                        "description": label
             })
        
        # Group entitiee
         {}
        for entity in entities:
            label = entity["label"]
            if label not in entit
                entity_groups[label] = []
            entity_groups[lab)
        
        return {
            "entities": entities,
            "groups": entity_groups,
        ntities)
        }n(e": letal_coun    "tot(entityel].appendoups:y_grs =y_groupentits by typ       es.appendRns.i patterint attern_lisabel, p for l   }b',\b'{2,4}\d1,2},?\s+s+\d{mber)\DeceNovember|r|Octobet|September|gusly|Au|Ju|Junepril|Mayry|March|A|Februa(?:January\br'2,4}\b[a-zZ][a-z]+\bN   |Foundationnstitutety|I(?:Universi\s&,.-]*ZA--Z][a-zr'\b[A       \.?)\by|LtdCompanration||Corpo|Corp\.?\.?|LLC?:Inc,.-]*("""gex patterns using reonextractiy st[Dict]](tllbacks_faitie_extract_entn self.  retur      ing fal us: {e},ction failedra
    
    def extract_dates_and_deadlines(self, text: str) -> List[Dict]:
        """Extract important dates and deadlines from legal text"""
        date_patterns = [
            r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b',  # MM/DD/YYYY or MM-DD-YYYY
            r'\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{2,4}\b',
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4}\b',
            r'\b\d{1,2}\s+days?\b',
            r'\b\d{1,2}\s+months?\b',
            r'\b\d{1,2}\s+years?\b'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Get context around the date
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()
                
                dates.append({
                    "date": match.group(),
                    "context": context,
                    "position": match.start()
                })
        
        return dates
    
    def extract_parties(self, text: str) -> List[str]:
        """Extract parties involved in the legal document"""
        # Common patterns for party identification
        party_patterns = [
            r'between\s+([^,\n]+)\s+and\s+([^,\n]+)',
            r'party\s+of\s+the\s+first\s+part[:\s]*([^,\n]+)',
            r'party\s+of\s+the\s+second\s+part[:\s]*([^,\n]+)',
            r'contracting\s+parties[:\s]*([^,\n]+)',
            r'(?:company|corporation|llc|inc\.?|ltd\.?)[:\s]*([^,\n]+)',
        ]
        
        parties = set()
        
        for pattern in party_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                for group in match.groups():
                    if group:
                        # Clean up the party name
                        party = re.sub(r'[^\w\s&,.-]', '', group.strip())
                        if len(party) > 2 and len(party) < 100:
                            parties.add(party)
        
        return list(parties)
    
    def process_document(self, file_path: str) -> Dict:
        """Complete document processing pipeline"""
        try:
            # Extract text
            raw_text = self.extract_text(file_path)
            if not raw_text:
                return {"error": "Could not extract text from document"}
            
            # Clean text
            clean_text = self.clean_text(raw_text)
            
            # Process document
            result = {
                "file_path": file_path,
                "raw_text": raw_text,
                "clean_text": clean_text,
                "word_count": len(clean_text.split()),
                "char_count": len(clean_text),
                "legal_sections": self.identify_legal_sections(clean_text),
                "entities": self.extract_entities(clean_text),
                "dates_and_deadlines": self.extract_dates_and_deadlines(clean_text),
                "parties": self.extract_parties(clean_text)
            }
            
            logger.info(f"Successfully processed document: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            return {"error": str(e)}

def main():
    """Test the document processor"""
    processor = DocumentProcessor()
    
    # Example usage
    test_text = """
    This Agreement is entered into between ABC Corporation, a Delaware corporation ("Company"), 
    and John Doe ("Employee") on January 1, 2024. The term of this agreement shall be 2 years.
    
    CONFIDENTIALITY: Employee agrees to maintain confidentiality of all proprietary information.
    
    TERMINATION: This agreement may be terminated by either party with 30 days notice.
    """
    
    # Test text processing
    clean_text = processor.clean_text(test_text)
    sections = processor.identify_legal_sections(clean_text)
    entities = processor.extract_entities(clean_text)
    dates = processor.extract_dates_and_deadlines(clean_text)
    parties = processor.extract_parties(clean_text)
    
    print("Legal Sections found:", list(sections.keys()))
    print("Entities found:", len(entities.get("entities", [])))
    print("Dates found:", len(dates))
    print("Parties found:", parties)

if __name__ == "__main__":
    main()
